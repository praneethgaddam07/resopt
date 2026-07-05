"""Standalone cover-letter generation — one heavy LLM call on résumé + JD + company.

Stateless and BYO-key. Unlike the résumé (which never names the hiring company),
a cover letter SHOULD address the company by name — that's expected and correct.
Grounded only in the candidate's real experience; never invents achievements.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt

from ..llm.client import LLMClient
from ..llm import prompts as P

_AI_TELL = re.compile(
    r"\b(delve|spearhead\w*|synerg\w*|orchestrat\w*|leverag\w*|harness\w*|foster\w*|"
    r"tapestry|unwavering|paramount|cutting-edge)\b", re.I)
_SYMBOLS = [("---", "-"), ("#", ""), ("_", " "), (";", ",")]


def generate_cover_letter(
    resume_text: str,
    jd_text: str,
    *,
    client: LLMClient,
    company: str,
    candidate_name: str,
    role: str = "",
    tone: str = "professional",
    hiring_manager: str = "",
) -> dict:
    """Return a structured cover letter {greeting, paragraphs[], closing, signature}."""
    ctx = f"MASTER RESUME:\n{resume_text}\n\nJOB DESCRIPTION:\n{jd_text}"
    details = (
        f"COMPANY: {company or '(the company)'}\n"
        f"ROLE: {role or '(infer the best-fit title from the JD)'}\n"
        f"CANDIDATE NAME: {candidate_name or 'the candidate'}\n"
        f"TONE: {tone or 'professional'}\n"
        f"ADDRESS TO: {hiring_manager.strip() or 'Hiring Team'}"
    )
    out = client.complete_json(
        ctx, P.COVER_LETTER + "\n\nDETAILS:\n" + details,
        mock=_mock(company, candidate_name, role),
        max_tokens=1200, task_tier="heavy")
    return _normalize(out, candidate_name, hiring_manager)


def _clean(s: str) -> str:
    s = str(s or "")
    s = _AI_TELL.sub("", s)
    for bad, good in _SYMBOLS:
        s = s.replace(bad, good)
    return re.sub(r"\s{2,}", " ", s).strip()


def _normalize(out: dict, candidate_name: str, hiring_manager: str) -> dict:
    out = dict(out or {})
    paras = out.get("paragraphs")
    if not isinstance(paras, list):
        paras = [str(paras)] if paras else []
    out["paragraphs"] = [_clean(p) for p in paras if str(p).strip()]
    greeting = str(out.get("greeting") or "").strip()
    if not greeting:
        greeting = f"Dear {hiring_manager.strip() or 'Hiring Team'},"
    out["greeting"] = greeting
    out["closing"] = str(out.get("closing") or "Sincerely,").strip()
    out["signature"] = (str(out.get("signature") or "").strip()
                        or candidate_name or "")
    return out


def cover_letter_text(letter: dict) -> str:
    """Flatten the structured letter into plain text (for display / clipboard)."""
    lines = [letter.get("greeting", "Dear Hiring Team,"), ""]
    for p in letter.get("paragraphs", []):
        lines += [p, ""]
    lines += [letter.get("closing", "Sincerely,"), letter.get("signature", "")]
    return "\n".join(l for l in lines).strip() + "\n"


def build_cover_letter_docx(letter: dict, candidate_name: str = "",
                            contact_line: str = "") -> bytes:
    """Render the letter as a clean, ATS-safe single-column .docx and return bytes."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Pt(54)
        section.top_margin = section.bottom_margin = Pt(54)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    name = letter.get("signature") or candidate_name
    if name:
        h = doc.add_paragraph()
        r = h.add_run(name)
        r.bold = True
        r.font.size = Pt(15)
        h.paragraph_format.space_after = Pt(0)
    if contact_line:
        c = doc.add_paragraph(contact_line)
        c.runs[0].font.size = Pt(10)
        c.paragraph_format.space_after = Pt(10)

    g = doc.add_paragraph(letter.get("greeting", "Dear Hiring Team,"))
    g.paragraph_format.space_after = Pt(8)
    for p in letter.get("paragraphs", []):
        para = doc.add_paragraph(p)
        para.paragraph_format.space_after = Pt(8)
    doc.add_paragraph(letter.get("closing", "Sincerely,")).paragraph_format.space_after = Pt(0)
    if name:
        doc.add_paragraph(name)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------- mock (no-key dev/tests) ---------------------------

def _mock(company: str, candidate_name: str, role: str) -> dict:
    co = company or "your team"
    title = role or "this role"
    return {
        "greeting": "Dear Hiring Team,",
        "paragraphs": [
            f"The {title} posting describes a team that needs faster, more reliable "
            f"reporting and someone who can own that problem end to end. That is exactly "
            f"the work I have done for the past three years.",
            "In my most recent role I reduced manual processing time by 34% and cut report "
            "turnaround from five days to under one, building automated validation that "
            "lifted data accuracy to 98%. Those are the same outcomes this role is asking for.",
            f"What draws me to {co} specifically is the chance to apply that same problem "
            "solving at a larger scale, turning scattered data work into decisions the whole "
            "team can rely on.",
            "I would welcome the chance to talk through how I can help. Thank you for your "
            "time and consideration.",
        ],
        "closing": "Sincerely,",
        "signature": candidate_name or "Candidate Name",
    }
