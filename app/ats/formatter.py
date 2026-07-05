"""Build an ATS-safe single-column .docx, applying portal rules (Step 15) + user style.

Universal rules enforced unconditionally regardless of style: single column (no
tables/columns), standard fonts only, contact in body text, no graphics. The
StyleOptions below are *cosmetic* overrides the user controls (font, per-section
sizes, alignment, bold, bullet glyph, spacing, margins, accent, education order,
fit-to-one-page) — none of them break ATS parsing.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, replace

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .portals import PortalRules, SECTION_LABELS

# ATS-safe font allowlist (ATS doc: standard fonts only — nothing decorative).
_SAFE_FONTS = {"Calibri", "Carlito", "Arial", "Helvetica", "Georgia",
               "Times New Roman", "Cambria", "Garamond"}
_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
          "right": WD_ALIGN_PARAGRAPH.RIGHT}


@dataclass
class StyleOptions:
    """User-controlled, ATS-safe formatting. Defaults reproduce the original output."""
    font: str = "Calibri"
    size_name: int = 16
    size_heading: int = 12          # SECTION headers (Summary, Skills, …)
    size_sub: int = 11              # role / company / degree lines
    size_body: int = 11
    header_align: str = "left"      # left | center | right (name + contact block)
    bold_headings: bool = True
    bold_sub: bool = True
    bullet: str = "•"               # "•" or "–"
    accent: str = ""                # hex without '#', "" = default near-black
    line_spacing: float = 1.0
    heading_gap: int = 4            # pt after a section heading
    entry_space: int = 2            # pt after bullets / sub lines
    section_space: int = 4          # pt gap after a whole entry block
    margin_v: int = 36              # top/bottom pt (36 = 0.5")
    margin_h: int = 36              # side pt
    education_order: str = "degree"  # degree | institution
    divider: bool = True            # subtle rule under section headings
    fit_one_page: bool = False
    skills_layout: str = "categorized"  # categorized | compact | inline (all ATS-safe, single column)
    date_format: str = "asis"       # asis | mon_year | num_year | year

    @classmethod
    def from_dict(cls, d: dict | None) -> "StyleOptions":
        """Build from an untrusted JSON dict, clamping/validating every field."""
        d = d or {}
        def _i(k, lo, hi, default):
            try:
                return max(lo, min(hi, int(round(float(d.get(k, default))))))
            except (TypeError, ValueError):
                return default
        font = str(d.get("font") or "Calibri")
        if font not in _SAFE_FONTS:
            font = "Calibri"
        align = str(d.get("header_align") or "left").lower()
        bullet = "–" if str(d.get("bullet") or "•") in ("–", "-", "dash") else "•"
        accent = "".join(c for c in str(d.get("accent") or "") if c in "0123456789abcdefABCDEF")
        accent = accent if len(accent) == 6 else ""
        try:
            ls = max(1.0, min(2.0, float(d.get("line_spacing", 1.0))))
        except (TypeError, ValueError):
            ls = 1.0
        return cls(
            font=font,
            size_name=_i("size_name", 11, 28, 16),
            size_heading=_i("size_heading", 9, 18, 12),
            size_sub=_i("size_sub", 9, 16, 11),
            size_body=_i("size_body", 8, 14, 11),
            header_align=align if align in _ALIGN else "left",
            bold_headings=bool(d.get("bold_headings", True)),
            bold_sub=bool(d.get("bold_sub", True)),
            bullet=bullet,
            accent=accent,
            line_spacing=ls,
            heading_gap=_i("heading_gap", 0, 12, 4),
            entry_space=_i("entry_space", 0, 10, 2),
            section_space=_i("section_space", 0, 16, 4),
            margin_v=_i("margin_v", 18, 90, 36),
            margin_h=_i("margin_h", 18, 90, 36),
            education_order=("institution" if str(d.get("education_order")) == "institution"
                             else "degree"),
            divider=bool(d.get("divider", True)),
            fit_one_page=bool(d.get("fit_one_page", False)),
            skills_layout=(str(d.get("skills_layout")) if str(d.get("skills_layout"))
                           in ("categorized", "compact", "inline") else "categorized"),
            date_format=(str(d.get("date_format")) if str(d.get("date_format"))
                         in ("asis", "mon_year", "num_year", "year") else "asis"),
        )


# --------------------- consistent spacing model ---------------------
# One hierarchy shared by BOTH the DOCX and PDF formatters so the two downloads
# space identically: bullets (entry_space) < entries (_entry_gap) < sections
# (_section_gap). Each section heading owns its top gap and each entry after the
# first owns its inter-entry gap, so spacing never depends on which sections or
# how many entries happen to precede it.

def _section_gap(st: "StyleOptions") -> int:
    """Space above every section heading — the largest gap in the hierarchy."""
    return st.section_space + 4


def _entry_gap(st: "StyleOptions") -> int:
    """Space above each entry (job/project/degree) after the first one in a section."""
    return st.section_space


# --------------------- low-level paragraph helpers ---------------------

def _para(doc, text, st: StyleOptions, *, bold=False, italic=False, size=None,
          space_after=None, space_before=None, caps=False, color="", align=None):
    p = doc.add_paragraph()
    run = p.add_run(text.upper() if caps else text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size if size is not None else st.size_body)
    run.font.name = st.font
    if color and len(color) == 6:
        run.font.color.rgb = RGBColor.from_string(color)
    pf = p.paragraph_format
    pf.space_after = Pt(st.entry_space if space_after is None else space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    pf.line_spacing = st.line_spacing
    if align in _ALIGN:
        p.alignment = _ALIGN[align]
    return p


def _content_width_emu(doc) -> int:
    s = doc.sections[0]
    return int(s.page_width - s.left_margin - s.right_margin)


def _two_col(doc, left, right, st: StyleOptions, *, bold=False, italic=False,
             size=None, space_after=None, space_before=None):
    """One line, `left` flush-left and `right` flush-right (single column, no table)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(st.entry_space if space_after is None else space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    pf.line_spacing = st.line_spacing
    pf.tab_stops.add_tab_stop(Emu(_content_width_emu(doc)), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(f"{left}\t{right}" if right else left)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size if size is not None else st.size_sub)
    run.font.name = st.font
    return p


def _looks_like_url(v: str) -> bool:
    """True only for values we can turn into a REAL, resolvable link — a URL or an
    email. A bare name/label ('LinkedIn', 'Venkata Gaddam') is NOT linkable and must
    render as plain text, or clicking it lands on a dead page."""
    v = (v or "").strip()
    if not v or " " in v.strip():          # multi-word => a label/name, not a URL
        return False
    if v.startswith(("http://", "https://", "mailto:")):
        return True
    if "@" in v and "/" not in v:          # email
        return True
    return bool(re.match(r"^[\w-]+(\.[\w-]+)+(/|$)", v))  # has a domain (a.b / a.b/c)


def _norm_url(v: str) -> str:
    """Add a scheme so the hyperlink actually resolves. Email -> mailto:."""
    v = (v or "").strip()
    if not v:
        return ""
    if v.startswith(("http://", "https://", "mailto:")):
        return v
    if "@" in v and "/" not in v:
        return "mailto:" + v
    return "https://" + v


def _add_hyperlink(paragraph, url: str, text: str, st: StyleOptions, size: int):
    """Append a real clickable hyperlink run (display text = `text`)."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi"):
        rf.set(qn(a), st.font)
    rPr.append(rf)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")  # standard hyperlink blue
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))  # half-points
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _render_contact_line(doc, contact: dict, st: StyleOptions, *, plain_urls: bool = False):
    """Build the contact line run-by-run so LinkedIn/GitHub/portfolio/email are
    clickable hyperlinks while phone/location stay plain text. `plain_urls=True`
    (iCIMS rule) renders them as visible plain text with no hyperlink at all."""
    size = max(9, st.size_body - 1)
    p = doc.add_paragraph()
    p.alignment = _ALIGN.get(st.header_align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = st.line_spacing
    first = True
    for k in ("phone", "email", "location", "linkedin", "github", "portfolio"):
        v = (contact.get(k) or "").strip()
        if not v:
            continue
        if not first:
            sep = p.add_run("    ")
            sep.font.size = Pt(size)
            sep.font.name = st.font
        first = False
        if k in ("linkedin", "github", "portfolio", "email") and not plain_urls \
                and _looks_like_url(v):
            _add_hyperlink(p, _norm_url(v), v, st, size)
        else:
            r = p.add_run(v)
            r.font.size = Pt(size)
            r.font.name = st.font


def _heading(doc, label, st: StyleOptions):
    # space_before gives every section a consistent top gap, independent of what
    # (or how much) preceded it — this is what makes section spacing even.
    p = _para(doc, label, st, bold=st.bold_headings, size=st.size_heading,
              space_after=st.heading_gap, space_before=_section_gap(st),
              caps=True, color=st.accent, align="left")
    if st.divider:
        _bottom_border(p)
    return p


def _bottom_border(p):
    """Subtle rule under a section heading (paragraph border — ATS-safe, not a table)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "2"), ("w:color", "BFBFBF")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


# --------------------- fit-to-one-page (best-effort heuristic) ---------------------

def _estimate_lines(data: dict, contact: dict) -> int:
    """Rough rendered-line estimate (~48 lines ≈ one page at default sizing)."""
    def wrapped(s, width=95):
        return 1 + max(0, len(s or "")) // width
    n = 2  # name + contact
    if data.get("summary"):
        n += 1 + wrapped(data["summary"], 90)
    skills = data.get("skills") or []
    if skills:
        n += 1 + sum(wrapped(", ".join(c.get("skills", []))) for c in skills)
    exps = data.get("experiences") or []
    if exps:
        n += 1
        for e in exps:
            n += 1 + (1 if e.get("bridge_line") else 0)
            n += sum(wrapped(b) for b in e.get("bullets", [])) + 1
    projs = data.get("projects") or []
    if projs:
        n += 1
        for p in projs:
            n += 1 + sum(wrapped(b) for b in p.get("bullets", []))
    edu = data.get("education") or contact.get("education") or []
    if edu:
        n += 1 + 2 * len(edu if isinstance(edu, list) else [edu])
    certs = data.get("certifications") or contact.get("certifications") or []
    if certs:
        n += 1 + len(certs)
    return n


def _fit(st: StyleOptions, est: int) -> StyleOptions:
    """Tiered compaction — scales sizes/spacing/margins down as content grows.
    Best-effort: a genuinely 2-page résumé cannot be losslessly forced to one page."""
    if est <= 48:
        return st
    if est <= 56:
        return replace(st, size_body=min(st.size_body, 10), size_sub=min(st.size_sub, 10),
                       heading_gap=3, entry_space=1, section_space=2, line_spacing=1.0,
                       margin_v=min(st.margin_v, 30), margin_h=min(st.margin_h, 32))
    if est <= 66:
        return replace(st, size_name=min(st.size_name, 15), size_heading=min(st.size_heading, 11),
                       size_sub=min(st.size_sub, 10), size_body=min(st.size_body, 10),
                       heading_gap=2, entry_space=1, section_space=2, line_spacing=1.0,
                       margin_v=27, margin_h=29)
    return replace(st, size_name=min(st.size_name, 14), size_heading=min(st.size_heading, 11),
                   size_sub=min(st.size_sub, 10), size_body=min(st.size_body, 9),
                   heading_gap=2, entry_space=0, section_space=1, line_spacing=1.0,
                   margin_v=22, margin_h=27)


# --------------------- main entry ---------------------

def build_docx_bytes(data: dict, contact: dict, portal: PortalRules,
                     section_order: list[str] | None = None,
                     style: StyleOptions | dict | None = None) -> bytes:
    """Render to .docx and return the bytes. `section_order` overrides the portal
    default (editor reorder); `style` carries the user's ATS-safe formatting."""
    st = style if isinstance(style, StyleOptions) else StyleOptions.from_dict(style)
    if st.date_format == "asis":
        # No explicit user choice -> impose the PORTAL's date style on 'Mon YYYY' tokens
        # (iCIMS: numeric 01/2024; others: abbreviated 'Jan 2024'). Never invents a day.
        st = replace(st, date_format=("num_year" if portal.date_format == "mmddyyyy"
                                      else "mon_year"))
    if st.fit_one_page:
        st = _fit(st, _estimate_lines(data, contact))

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Pt(st.margin_h)
        section.top_margin = section.bottom_margin = Pt(st.margin_v)

    normal = doc.styles["Normal"]
    normal.font.name = st.font
    normal.font.size = Pt(st.size_body)

    education = data.get("education") or contact.get("education") or []
    certs = data.get("certifications") or contact.get("certifications") or []

    renderers = {
        "contact": lambda: _render_contact(doc, data, contact, portal, st),
        "summary": lambda: _render_summary(doc, data, portal, st),
        "skills": lambda: _render_skills(doc, data, st),
        "certifications": lambda: _render_certs(doc, certs, portal, st),
        "experience": lambda: _render_experience(doc, data, portal, st),
        "projects": lambda: _render_projects(doc, data, st),
        "education": lambda: _render_education(doc, education, st),
    }
    seen = set()
    for key in (section_order or portal.section_order):
        if key in renderers and key not in seen:
            seen.add(key)
            renderers[key]()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx(data: dict, contact: dict, portal: PortalRules, out_path: str,
               style: StyleOptions | dict | None = None) -> None:
    with open(out_path, "wb") as f:
        f.write(build_docx_bytes(data, contact, portal, style=style))


# --------------------- section renderers ---------------------

def _render_contact(doc, data, contact, portal: PortalRules, st: StyleOptions):
    p = doc.add_paragraph()
    run = p.add_run(contact.get("name", "Candidate Name"))
    run.bold = True
    run.font.size = Pt(st.size_name)
    run.font.name = st.font
    p.alignment = _ALIGN.get(st.header_align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.line_spacing = st.line_spacing

    # Contact line — email/LinkedIn/GitHub/portfolio render as clickable hyperlinks
    # (display text = the full URL, so ATS still reads it; humans can click).
    # iCIMS demands full visible plain text, NOT hyperlinked — honor the portal flag.
    _render_contact_line(doc, contact, st, plain_urls=portal.plain_text_urls)

    if portal.job_title_tagline_page1 and data.get("job_title"):  # Taleo
        _para(doc, data["job_title"], st, bold=True, size=st.size_heading, space_after=6,
              align=st.header_align)


def _render_summary(doc, data, portal: PortalRules, st: StyleOptions):
    if data.get("job_title") and not portal.job_title_tagline_page1:
        _para(doc, data["job_title"], st, bold=st.bold_sub, size=st.size_heading,
              space_after=st.heading_gap, align=st.header_align)
    _heading(doc, SECTION_LABELS["summary"], st)
    # Trailing gap is owned by the NEXT section's heading, so keep this tight.
    _para(doc, data.get("summary", ""), st, space_after=st.entry_space)


_MONTHS_FULL = ["January", "February", "March", "April", "May", "June",
                "July", "August", "September", "October", "November", "December"]
_MONTH_IDX = {}
for _i, _m in enumerate(_MONTHS_FULL, 1):
    _MONTH_IDX[_m.lower()] = _i
    _MONTH_IDX[_m[:3].lower()] = _i


def _fmt_one_date(tok: str, fmt: str) -> str:
    """Reformat one 'Month YYYY' token; leave anything else (Present, bare year) as-is."""
    m = re.match(r"^([A-Za-z]{3,9})\.?\s+(\d{4})$", tok.strip())
    if not m:
        return tok.strip()
    idx = _MONTH_IDX.get(m.group(1).lower())
    if not idx:
        return tok.strip()
    yr = m.group(2)
    if fmt == "num_year":
        return f"{idx:02d}/{yr}"
    if fmt == "year":
        return yr
    return f"{_MONTHS_FULL[idx - 1][:3]} {yr}"  # mon_year -> abbreviated month


def _fmt_dates(s: str, fmt: str) -> str:
    """Best-effort reformat of a date or range, e.g. 'May 2023 - Aug 2024'. ATS-safe text."""
    if not s or fmt == "asis":
        return s
    parts = re.split(r"(\s*[-–—]\s*|\s+to\s+)", s)
    return "".join(_fmt_one_date(p, fmt) if i % 2 == 0 else p for i, p in enumerate(parts))


def _render_skills(doc, data, st: StyleOptions):
    _heading(doc, SECTION_LABELS["skills"], st)
    cats = [c for c in data.get("skills", []) if c.get("skills")]
    # Inline: one dense paragraph, no category labels — the most compact, still single column.
    if st.skills_layout == "inline":
        flat = [s for c in cats for s in c.get("skills", [])]
        if flat:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(st.entry_space + 1)
            p.paragraph_format.line_spacing = st.line_spacing
            r = p.add_run(", ".join(flat) + ".")
            r.font.size = Pt(st.size_body)
            r.font.name = st.font
        return
    # Categorized (default) or Compact (same, but no gap between category lines).
    gap = 0 if st.skills_layout == "compact" else st.entry_space + 1
    for cat in cats:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(gap)
        p.paragraph_format.line_spacing = st.line_spacing
        head = p.add_run(f"{cat.get('name','')}: ")
        head.bold = st.bold_sub
        head.font.size = Pt(st.size_body)
        head.font.name = st.font
        body = p.add_run(", ".join(cat.get("skills", [])) + ".")
        body.font.size = Pt(st.size_body)
        body.font.name = st.font


def _render_certs(doc, certs, portal: PortalRules, st: StyleOptions):
    if not certs:
        return
    _heading(doc, SECTION_LABELS["certifications"], st)
    for c in certs:
        if isinstance(c, dict):
            line = c.get("name", "")
            if c.get("issuer"):
                line += f" — {c['issuer']}"
            if c.get("url"):
                line += f"  {c['url']}"
            _para(doc, line, st)
        else:
            _para(doc, str(c), st)


def _render_experience(doc, data, portal: PortalRules, st: StyleOptions):
    _heading(doc, SECTION_LABELS["experience"], st)
    for i, exp in enumerate(data.get("experiences", [])):
        left = ", ".join(x for x in [exp.get("title", ""), exp.get("company", "")] if x)
        right = " | ".join(x for x in [_fmt_dates(exp.get("duration", ""), st.date_format),
                                       exp.get("location", "")] if x)
        # Even inter-entry gap on the role header (2nd job onward); no empty spacer paragraph.
        _two_col(doc, left, right, st, bold=st.bold_sub, space_after=1,
                 space_before=(_entry_gap(st) if i else None))
        if exp.get("bridge_line"):
            _para(doc, exp["bridge_line"], st, italic=True, size=max(9, st.size_body - 1))
        for b in exp.get("bullets", []):
            _para(doc, f"{st.bullet} {b}", st)


def _render_projects(doc, data, st: StyleOptions):
    projects = data.get("projects", [])
    if not projects:
        return
    _heading(doc, SECTION_LABELS["projects"], st)
    for i, proj in enumerate(projects):
        # Same even inter-entry gap as experience (was missing entirely before).
        _para(doc, proj.get("title", "Project"), st, bold=st.bold_sub, space_after=1,
              space_before=(_entry_gap(st) if i else None))
        for b in proj.get("bullets", []):
            _para(doc, f"{st.bullet} {b}", st)


def _render_education(doc, education, st: StyleOptions):
    if not education:
        return
    _heading(doc, SECTION_LABELS["education"], st)
    rendered = 0  # count only entries actually emitted, so the FIRST gets no top gap
    for item in (education if isinstance(education, list) else [education]):
        if isinstance(item, dict):
            degree = item.get("degree", "").strip()
            school = item.get("school", "").strip()
            dates = _fmt_dates(item.get("dates", "").strip(), st.date_format)
            gpa = item.get("gpa", "").strip()
            if not (degree or school):
                continue
            # "Show education by": which line leads (degree-first vs institution-first).
            primary, secondary = (degree, school) if st.education_order == "degree" else (school, degree)
            _two_col(doc, primary, dates, st, bold=st.bold_sub, space_after=0,
                     space_before=(_entry_gap(st) if rendered else None))
            tail = secondary + (f", GPA: {gpa}" if gpa else "")
            if tail:
                _para(doc, tail, st, space_after=st.entry_space)
            rendered += 1
        else:
            s = (item or "").strip()
            if s:
                _para(doc, s, st, space_before=(_entry_gap(st) if rendered else None))
                rendered += 1
