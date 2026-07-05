"""ATS portal rules + DOCX formatting."""
import io

from docx import Document

from app.ats.portals import get_portal, list_portals
from app.ats.formatter import build_docx_bytes

DATA = {
    "job_title": "Data Analyst",
    "summary": "Data Analyst who delivered measurable reporting reliability gains.",
    "skills": [{"name": "Core", "skills": ["SQL", "Python"]},
               {"name": "BI", "skills": ["Tableau"]}],
    "experiences": [{
        "title": "Most Recent Role", "company": "Confidential",
        "duration": "May 2022 – Present", "location": "Remote",
        "bridge_line": "Maps directly to the team's reporting reliability problem.",
        "bullets": ["Built dashboards using SQL that reduced reporting time by 34%."],
    }],
    "projects": [{"title": "Project 1", "bullets": ["Automated validation improving accuracy to 98%."]}],
    "education": "Master of Science in Information Technology, State University",
    "certifications": [{"name": "AWS Certified", "issuer": "Amazon Web Services", "url": "verify.example/abc"}],
}
CONTACT = {"name": "Jane Doe", "phone": "555", "email": "j@x.com",
           "linkedin": "linkedin.com/in/janedoe", "location": "NY",
           "education": DATA["education"], "certifications": DATA["certifications"]}


def _headings(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    # All-caps bold-ish lines are our section headings.
    return [p.text for p in doc.paragraphs if p.text.isupper() and p.text.strip()]


def test_all_portals_render():
    for p in list_portals():
        portal = get_portal(p["key"])
        b = build_docx_bytes(DATA, CONTACT, portal)
        assert b[:2] == b"PK"  # valid .docx (zip) header


def test_spacing_is_even_across_entries_and_sections():
    """Regression: exports must space uniformly — no empty spacer paragraphs, one
    consistent gap above every section heading, and one consistent gap between
    entries (jobs jammed differently from projects was the reported bug)."""
    from app.ats.formatter import StyleOptions, _section_gap, _entry_gap
    data = dict(DATA)
    data["experiences"] = [
        {"title": "Role A", "company": "Co A", "duration": "May 2022 - Present",
         "location": "Remote", "bullets": ["Did A improving x by 20%.", "Did B."]},
        {"title": "Role B", "company": "Co B", "duration": "Jan 2020 - May 2022",
         "location": "NY", "bullets": ["Did C.", "Did D reducing y by 30%."]},
    ]
    data["projects"] = [
        {"title": "Proj 1", "bullets": ["Built p1."]},
        {"title": "Proj 2", "bullets": ["Built p2."]},
    ]
    data["education"] = [
        {"degree": "Master of Science", "school": "State U", "dates": "2022 - 2024", "gpa": "3.6"},
        {"degree": "Bachelor of Science", "school": "Tech U", "dates": "2018 - 2022", "gpa": "3.5"},
    ]
    st = StyleOptions()
    doc = Document(io.BytesIO(build_docx_bytes(data, CONTACT, get_portal("generic"), style=st)))

    # 1) no empty spacer paragraphs
    assert all(p.text.strip() for p in doc.paragraphs), "empty spacer paragraph leaked in"

    befores = [p.paragraph_format.space_before.pt for p in doc.paragraphs
               if p.paragraph_format.space_before is not None]
    # 2) exactly two consistent gap values exist: section gap and entry gap
    assert set(round(b) for b in befores) == {_section_gap(st), _entry_gap(st)}
    # 3) every section heading uses the (larger) section gap
    heads = [p for p in doc.paragraphs if p.text.isupper() and p.text.strip()]
    assert heads and all(round(p.paragraph_format.space_before.pt) == _section_gap(st)
                         for p in heads)
    # 4) section gap is strictly larger than the entry gap (correct hierarchy)
    assert _section_gap(st) > _entry_gap(st) > 0


def test_pdf_vspace_treats_zero_as_zero():
    """Root cause of the PDF education-gap bug: fpdf2's ln(0) falls back to the
    last cell height (a phantom blank line). _vspace must move exactly zero."""
    from app.ats.pdf_formatter import _Pdf
    from app.ats.formatter import StyleOptions
    from fpdf.enums import XPos, YPos
    pdf = _Pdf(StyleOptions())
    pdf.set_font(pdf.fam, "", 11)
    pdf.cell(100, pdf._lh(11), "x", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    y = pdf.get_y()
    pdf._vspace(0)
    assert pdf.get_y() == y            # zero gap => no movement (not a phantom line)
    pdf._vspace(5)
    assert round(pdf.get_y() - y, 1) == 5.0


def test_pdf_education_degree_line_no_phantom_gap():
    """Regression: an education entry's degree line (two_col, space_after=0) must
    advance exactly ONE line to its school line — not two (the reported PDF bug
    where every degree was pushed far from its university)."""
    from app.ats.pdf_formatter import _Pdf
    from app.ats.formatter import StyleOptions
    st = StyleOptions()
    pdf = _Pdf(st)
    y0 = pdf.get_y()
    pdf.two_col("Master of Science, Information Technology and Management",
                "Aug 2024 - May 2026", bold=st.bold_sub, space_after=0)
    assert round(pdf.get_y() - y0, 1) == round(pdf._lh(st.size_sub), 1)


def test_adp_section_order():
    portal = get_portal("adp")
    heads = _headings(build_docx_bytes(DATA, CONTACT, portal))
    # ADP mandate: Summary -> Skills -> Experience -> Education -> Certifications.
    assert heads.index("SUMMARY") < heads.index("SKILLS") < heads.index("WORK EXPERIENCE")
    assert heads.index("EDUCATION") < heads.index("CERTIFICATIONS")


def test_amazon_certs_before_skills():
    portal = get_portal("amazon")
    heads = _headings(build_docx_bytes(DATA, CONTACT, portal))
    assert heads.index("CERTIFICATIONS") < heads.index("SKILLS")


def test_date_formats():
    assert get_portal("icims").date_example() == "05/01/2023"
    assert "–" in get_portal("workday").date_example()  # em-dash range



def test_skills_layout_and_date_format():
    """Skills layout (inline collapses categories) + best-effort date reformatting."""
    from io import BytesIO
    from docx import Document
    from app.ats.formatter import build_docx_bytes, StyleOptions
    from app.ats.portals import get_portal
    data = {"job_title": "Data Analyst", "summary": "S",
            "skills": [{"name": "Cat A", "skills": ["Python", "SQL"]},
                       {"name": "Cat B", "skills": ["Tableau", "AWS"]}],
            "experiences": [{"title": "Analyst", "company": "Co",
                             "duration": "May 2023 - Aug 2024", "location": "X",
                             "bullets": ["did x"]}],
            "projects": [],
            "education": [{"degree": "BS", "school": "U", "dates": "Aug 2020 - May 2024", "gpa": ""}],
            "certifications": []}
    contact, portal = {"name": "J"}, get_portal("generic")

    def text(st):
        b = build_docx_bytes(data, contact, portal, style=st)
        return "\n".join(pr.text for pr in Document(BytesIO(b)).paragraphs)

    inline = text(StyleOptions(skills_layout="inline"))
    assert "Python, SQL, Tableau, AWS." in inline and "Cat A:" not in inline
    assert "Cat A:" in text(StyleOptions(skills_layout="categorized"))
    assert "05/2023" in text(StyleOptions(date_format="num_year"))
    assert "2023 - 2024" in text(StyleOptions(date_format="year"))


def _document_xml(docx_bytes):
    import zipfile
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        return z.read("word/document.xml").decode("utf-8")


def test_icims_urls_are_plain_text_not_hyperlinked():
    # iCIMS rule: URLs as full visible plain text — the portal flag must kill hyperlinks.
    icims = _document_xml(build_docx_bytes(DATA, CONTACT, get_portal("icims")))
    generic = _document_xml(build_docx_bytes(DATA, CONTACT, get_portal("generic")))
    assert "<w:hyperlink" not in icims
    assert "<w:hyperlink" in generic          # everyone else keeps clickable links
    assert "linkedin.com/in/janedoe" in icims  # the URL text itself still present


def test_portal_date_format_applied_by_default():
    data = dict(DATA)
    data["experiences"] = [dict(DATA["experiences"][0], duration="January 2023 - August 2024")]
    icims = _document_xml(build_docx_bytes(data, CONTACT, get_portal("icims")))
    assert "01/2023" in icims and "08/2024" in icims      # numeric for iCIMS
    generic = _document_xml(build_docx_bytes(data, CONTACT, get_portal("generic")))
    assert "Jan 2023" in generic and "Aug 2024" in generic  # abbreviated elsewhere
    # An explicit user choice still wins over the portal default.
    keep = _document_xml(build_docx_bytes(data, CONTACT, get_portal("icims"),
                                          style={"date_format": "mon_year"}))
    assert "Jan 2023" in keep


def test_pdf_renders_all_portals_and_mirrors_rules():
    from app.ats.pdf_formatter import build_pdf_bytes
    for p in list_portals():
        b = build_pdf_bytes(DATA, CONTACT, get_portal(p["key"]))
        assert b[:5] == b"%PDF-", f"{p['key']} produced invalid PDF"
    # unicode bullets/dashes must not crash the latin-1 core fonts
    data = dict(DATA)
    data["experiences"] = [dict(DATA["experiences"][0],
                                duration="January 2023 – August 2024",
                                bullets=["Built • dashboards — cutting résumé time by 34%"])]
    b = build_pdf_bytes(data, CONTACT, get_portal("icims"))
    assert b[:5] == b"%PDF-"


def test_only_real_urls_get_hyperlinked_not_names():
    from app.ats.formatter import _looks_like_url
    # real, linkable
    assert _looks_like_url("linkedin.com/in/janedoe")
    assert _looks_like_url("https://github.com/x")
    assert _looks_like_url("jane@x.com")
    # NOT linkable — bare labels / names (the dead-link bug)
    assert not _looks_like_url("Venkata Praneeth Gaddam")
    assert not _looks_like_url("LinkedIn")
    assert not _looks_like_url("")
    # a name in the linkedin field must NOT become a <w:hyperlink>
    bad = dict(CONTACT, linkedin="Venkata Praneeth Gaddam")
    xml = _document_xml(build_docx_bytes(DATA, bad, get_portal("generic")))
    assert "Venkata Praneeth Gaddam" in xml and "<w:hyperlink" not in xml.split("Venkata")[0][-200:]


def test_extract_reads_urls_behind_hyperlinks():
    import io as _io
    from docx import Document as _Doc
    from app.workflow.extract import extract_text_from_file
    doc = _Doc()
    p = doc.add_paragraph()
    # display text "LinkedIn" but the real target is the URL (the exact failure case)
    part = p.part
    r_id = part.relate_to("https://www.linkedin.com/in/praneeth-real",
                          "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "LinkedIn"
    run.append(t); hl.append(run); p._p.append(hl)
    buf = _io.BytesIO(); doc.save(buf)
    text = extract_text_from_file("r.docx", buf.getvalue())
    assert "linkedin.com/in/praneeth-real" in text  # the REAL url now reaches extraction
